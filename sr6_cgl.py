import sys
import csv
from docx import Document
from docx.shared import Pt
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import RGBColor
from docx.shared import Inches

CGL_NORMAL_TEXT = 'Normal Text'
CGL_SHADOWTALK_STYLE_NAME = 'Shadowtalk'
CGL_LAYOUT_NOTE_STYLE_NAME = 'Layout Note'
CGL_LAYOUT_HEADER_ONE_STYLE_NAME = "Header 1"
CGL_LAYOUT_HEADER_TWO_STYLE_NAME = "Header 2"
CGL_LAYOUT_HEADER_THREE_STYLE_NAME = "Header 3"
CGL_LAYOUT_HEADER_FOUR_STYLE_NAME = "Header 4"
CGL_LAYOUT_HEADER_FIVE_STYLE_NAME = "Header 5"

DEBUG = False

def openDocument(filename):
    f = open(input_filename, 'rb')
    document = Document(f)
    f.close()
    return document

def delete_paragraph(paragraph):
    p = paragraph._element
    p.getparent().remove(p)
    p._p = p._element = None

def shadowtalk(paragraph, document):
    paragraph.insert_paragraph_before(paragraph.text.replace(">",">"), document.styles[CGL_SHADOWTALK_STYLE_NAME])
    delete_paragraph(paragraph)

def shadowtalks(document):
    st_begin = False
    for paragraph in document.paragraphs:
        if paragraph.text.startswith(">") and not paragraph.text.startswith(">>>>>"):
            st_begin = not st_begin
            shadowtalk(paragraph, document)
        else:
            if st_begin:
                shadowtalk(paragraph, document)

def layout_note(paragraph, text, document, style):
    paragraph.insert_paragraph_before(text,style)
    delete_paragraph(paragraph)

def layout_notes(document):
    layout_note_begin = False
    nb_layout_notes = 1
    for paragraph in document.paragraphs:
        if paragraph.text.startswith(">>>"):    # if there is more than one > it's most likely a layout note
            if not layout_note_begin:
                if ( DEBUG ): print("Layout notes detected [%i]" % nb_layout_notes)
                nb_layout_notes += 1
            layout_note_begin = not layout_note_begin
            layout_note(paragraph,paragraph.text.replace(">>>>>","/////"), document, document.styles[CGL_LAYOUT_NOTE_STYLE_NAME])
        if paragraph.text.startswith(CGL_LAYOUT_NOTE_STYLE_NAME):
            layout_note_begin = not layout_note_begin

    for paragraph in document.paragraphs:
        if paragraph.text.startswith(">>>>>"):
            delete_paragraph(paragraph)

    for paragraph in document.paragraphs:
        if paragraph.text.startswith("/////"):
            paragraph.text.replace("/////",">>>>>")

def map_gdoc_styles_to_sr6(document):
    for paragraph in document.paragraphs:
        if paragraph.style.name in index.keys():
            paragraph.style.name = index[paragraph.style.name]


if len(sys.argv) != 3:
    raise ValueError("Invalid number of parameters")

input_filename = sys.argv[1]
output_filename = sys.argv[2]

index = {
	"LO-normal": CGL_NORMAL_TEXT,
    "normal": CGL_NORMAL_TEXT,
    "normal1": CGL_NORMAL_TEXT,
	"Heading 1": CGL_LAYOUT_HEADER_TWO_STYLE_NAME,
	"Heading 2": CGL_LAYOUT_HEADER_THREE_STYLE_NAME,
	"Heading 3": CGL_LAYOUT_HEADER_FOUR_STYLE_NAME,
    "Heading 4": CGL_LAYOUT_HEADER_FIVE_STYLE_NAME,
    "Title": CGL_LAYOUT_HEADER_ONE_STYLE_NAME,
}

document = openDocument(input_filename)

document.styles.add_style(CGL_LAYOUT_HEADER_ONE_STYLE_NAME, WD_STYLE_TYPE.PARAGRAPH)
document.styles[CGL_LAYOUT_HEADER_ONE_STYLE_NAME].font.bold = True
document.styles[CGL_LAYOUT_HEADER_ONE_STYLE_NAME].font.name = 'Times New Roman'
document.styles[CGL_LAYOUT_HEADER_ONE_STYLE_NAME].font.size = Pt(24)
document.styles[CGL_LAYOUT_HEADER_ONE_STYLE_NAME].font.small_caps = True

document.styles.add_style(CGL_LAYOUT_HEADER_TWO_STYLE_NAME, WD_STYLE_TYPE.PARAGRAPH)
document.styles[CGL_LAYOUT_HEADER_TWO_STYLE_NAME].font.bold = True
document.styles[CGL_LAYOUT_HEADER_TWO_STYLE_NAME].font.name = 'Times New Roman'
document.styles[CGL_LAYOUT_HEADER_TWO_STYLE_NAME].font.size = Pt(14)
document.styles[CGL_LAYOUT_HEADER_TWO_STYLE_NAME].font.small_caps = True

document.styles.add_style(CGL_LAYOUT_HEADER_THREE_STYLE_NAME, WD_STYLE_TYPE.PARAGRAPH)
document.styles[CGL_LAYOUT_HEADER_THREE_STYLE_NAME].font.bold = True
document.styles[CGL_LAYOUT_HEADER_THREE_STYLE_NAME].font.name = 'Times New Roman'
document.styles[CGL_LAYOUT_HEADER_THREE_STYLE_NAME].font.size = Pt(10)
document.styles[CGL_LAYOUT_HEADER_THREE_STYLE_NAME].font.small_caps = True

document.styles.add_style(CGL_LAYOUT_HEADER_FOUR_STYLE_NAME, WD_STYLE_TYPE.PARAGRAPH)
document.styles[CGL_LAYOUT_HEADER_FOUR_STYLE_NAME].font.bold = True
document.styles[CGL_LAYOUT_HEADER_FOUR_STYLE_NAME].font.name = 'Times New Roman'
document.styles[CGL_LAYOUT_HEADER_FOUR_STYLE_NAME].font.size = Pt(9)

document.styles.add_style(CGL_LAYOUT_HEADER_FIVE_STYLE_NAME, WD_STYLE_TYPE.PARAGRAPH)
document.styles[CGL_LAYOUT_HEADER_FIVE_STYLE_NAME].font.italic = True
document.styles[CGL_LAYOUT_HEADER_FIVE_STYLE_NAME].font.name = 'Times New Roman'
document.styles[CGL_LAYOUT_HEADER_FIVE_STYLE_NAME].font.size = Pt(9)

document.styles.add_style(CGL_NORMAL_TEXT, WD_STYLE_TYPE.PARAGRAPH)
document.styles[CGL_NORMAL_TEXT].font.italic = False
document.styles[CGL_NORMAL_TEXT].font.name = 'Times New Roman'
document.styles[CGL_NORMAL_TEXT].font.size = Pt(9)

document.styles.add_style(CGL_SHADOWTALK_STYLE_NAME, WD_STYLE_TYPE.PARAGRAPH)
document.styles[CGL_SHADOWTALK_STYLE_NAME].font.italic = False
document.styles[CGL_SHADOWTALK_STYLE_NAME].font.name = 'Verdana'
document.styles[CGL_SHADOWTALK_STYLE_NAME].font.size = Pt(9)

document.styles.add_style(CGL_LAYOUT_NOTE_STYLE_NAME, WD_STYLE_TYPE.PARAGRAPH)
document.styles[CGL_LAYOUT_NOTE_STYLE_NAME].font.color.rgb = RGBColor(255, 0, 0)
document.styles[CGL_LAYOUT_NOTE_STYLE_NAME].font.name = 'Times New Roman'
document.styles[CGL_LAYOUT_NOTE_STYLE_NAME].font.size = Pt(12)

map_gdoc_styles_to_sr6(document)
shadowtalks(document)
layout_notes(document)
document.save(output_filename)
