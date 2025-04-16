import sys
import csv
import re
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Inches

# TODO: liste à puces
BBE_SR5_TEXT = 'SR5 - Maquette'
BBE_SR6_TEXT = "SR6 - Texte"
BBE_SR6_TITLE_1 = "SR6 - Titre 1"
BBE_SR6_TITLE_2 = "SR6 - Titre 2"
BBE_SR6_TITLE_3 = "SR6 - Titre 3"
BBE_SR6_TITLE_4 = "SR6 - Titre 4"
BBE_SR6_TITLE_CHAPTER = "SR6 - Titre chapitre"
BBE_SR6_INSERT_TEXT = 'SR6 - Encart texte'
BBE_SR6_INSERT_TITLE_1 = 'SR6 - Encart titre 1'
BBE_SR6_INSERT_LIST = 'SR6 - Encart texte (liste)'
BBE_SR6_INSERT_TITLE_2 = 'SR6 - Encart titre 2'
BBE_SR6_STATS_DROP_WEAPONS = 'SR6 - Stat décroché armes'
BBE_SR6_STATS_DROP = 'SR6 - Stat décroché (bold :)'
BBE_SR6_STATS_TITLE_1 = 'SR6 - Stat titre 1'
BBE_SR6_STATS_TITLE_2 = 'SR6 - Stat titre 2'
BBE_SR6_TABLE_HEADER = 'SR6 - Table en-tête'
BBE_SR6_TABLE_TEXT = 'SR6 - Table texte'
BBE_SR6_SHADOWTALK = 'SR6 - Shadowtalk'
BBE_SR6_POSTED_BY = 'SR6 - Posté par'

def openDocument(filename):
    f = open(input_filename, 'rb')
    document = Document(f)
    f.close()
    return document

def delete_paragraph(paragraph):
    p = paragraph._element
    p.getparent().remove(p)
    p._p = p._element = None

def shadowtalk(paragraph, document):
    paragraph.insert_paragraph_before(paragraph.text.replace("> ",""),document.styles["SR6 - Shadowtalk"])
    delete_paragraph(paragraph)


def shadowtalks(document):
    st_begin = False
    for paragraph in document.paragraphs:
        if paragraph.text.startswith(">"):
            st_begin = not st_begin
            shadowtalk(paragraph, document)
        else:
            if st_begin:
                shadowtalk(paragraph, document)

def encart(paragraph, text, document, style):
    paragraph.insert_paragraph_before(text,style)
    delete_paragraph(paragraph)

def encarts(document):
    encart_begin = False
    for paragraph in document.paragraphs:
        if paragraph.text.startswith("Encart:") or paragraph.text.startswith("Encart :"):
            print(paragraph.text)
            encart_begin = not encart_begin
            encart(paragraph,re.sub(r"^Encart\s?:\s?", "", paragraph.text), document, document.styles['SR6 - Encart titre 1'])
        else:
            if encart_begin:
                encart(paragraph, paragraph.text, document, document.styles['SR6 - Encart texte'])
        if paragraph.text.startswith("Fin Encart"):
            #encart(paragraph, " ", document, document.styles['SR6 - Encart texte'])
            print(paragraph.text)
            encart_begin = not encart_begin

    for paragraph in document.paragraphs:
        if paragraph.text.startswith("Fin Encart"):
            delete_paragraph(paragraph)

def map_gdoc_styles_to_sr6(document):
    for paragraph in document.paragraphs:
        if paragraph.style.name in index.keys():
            paragraph.style.name = index[paragraph.style.name]



if len(sys.argv) != 3:
    raise ValueError("Invalid number of parameters")

input_filename = sys.argv[1]
output_filename = sys.argv[2]

index = {
    "LO-normal": BBE_SR6_TEXT,
    "Normal": BBE_SR6_TEXT,
    "normal": BBE_SR6_TEXT,
    "normal1": BBE_SR6_TEXT,
    "Heading 1": BBE_SR6_TITLE_1,
    "Heading 2": BBE_SR6_TITLE_2,
    "Heading 3": BBE_SR6_TITLE_3,
    "Heading 4": BBE_SR6_TITLE_4,
    "Title": BBE_SR6_TITLE_CHAPTER,
}

document = openDocument(input_filename)
document.styles.add_style(BBE_SR5_TEXT, WD_STYLE_TYPE.PARAGRAPH)
document.styles.add_style(BBE_SR6_TEXT, WD_STYLE_TYPE.PARAGRAPH)
document.styles.add_style(BBE_SR6_TITLE_1, WD_STYLE_TYPE.PARAGRAPH)
document.styles.add_style(BBE_SR6_TITLE_2, WD_STYLE_TYPE.PARAGRAPH)
document.styles.add_style(BBE_SR6_TITLE_3, WD_STYLE_TYPE.PARAGRAPH)
document.styles.add_style(BBE_SR6_TITLE_4, WD_STYLE_TYPE.PARAGRAPH)
document.styles.add_style(BBE_SR6_TITLE_CHAPTER, WD_STYLE_TYPE.PARAGRAPH)
document.styles.add_style(BBE_SR6_INSERT_TEXT, WD_STYLE_TYPE.PARAGRAPH).paragraph_format.left_indent = Inches(0.5)
document.styles.add_style(BBE_SR6_INSERT_TITLE_1, WD_STYLE_TYPE.PARAGRAPH).paragraph_format.left_indent = Inches(0.5)
document.styles.add_style(BBE_SR6_INSERT_LIST, WD_STYLE_TYPE.PARAGRAPH)
document.styles.add_style(BBE_SR6_INSERT_TITLE_2, WD_STYLE_TYPE.PARAGRAPH)
document.styles.add_style(BBE_SR6_STATS_DROP_WEAPONS, WD_STYLE_TYPE.PARAGRAPH)
document.styles.add_style(BBE_SR6_STATS_DROP, WD_STYLE_TYPE.PARAGRAPH)
document.styles.add_style(BBE_SR6_STATS_TITLE_1, WD_STYLE_TYPE.PARAGRAPH)
document.styles.add_style(BBE_SR6_STATS_TITLE_2, WD_STYLE_TYPE.PARAGRAPH)
document.styles.add_style(BBE_SR6_TABLE_HEADER, WD_STYLE_TYPE.PARAGRAPH)
document.styles.add_style(BBE_SR6_TABLE_TEXT, WD_STYLE_TYPE.PARAGRAPH)
document.styles.add_style(BBE_SR6_SHADOWTALK, WD_STYLE_TYPE.PARAGRAPH)
document.styles.add_style(BBE_SR6_POSTED_BY, WD_STYLE_TYPE.PARAGRAPH)
document.styles[BBE_SR6_SHADOWTALK].font.italic = True

map_gdoc_styles_to_sr6(document)

for paragraph in document.paragraphs:
    if paragraph.text.startswith("Posté par"):
        paragraph.insert_paragraph_before(paragraph.text,document.styles["SR6 - Posté par"])
        delete_paragraph(paragraph)

shadowtalks(document)
encarts(document)
document.save(output_filename)
